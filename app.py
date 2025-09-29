import streamlit as st
from fpdf import FPDF
from docx import Document
from datetime import datetime
import os

st.set_page_config(page_title="Betriebsanweisung Generator", layout="centered")

st.title("📄 Betriebsanweisung Generator nach DGUV 211-010")

# Eingabefelder
firmenname = st.text_input("Firmenname")
nr = st.text_input("Nummer")
arbeitsbereich = st.text_input("Arbeitsbereich")
arbeitsplatz = st.text_input("Arbeitsplatz")
taetigkeit = st.text_input("Tätigkeit")

gefahren = st.text_area("Gefahren für Mensch und Umwelt")
schutz = st.text_area("Schutzmaßnahmen und Verhaltensregeln")
verhalten_gefahrfall = st.text_area("Verhalten im Gefahrfall")
erste_hilfe = st.text_area("Erste Hilfe")
entsorgung = st.text_area("Sachgerechte Entsorgung")

pictogramme = st.file_uploader("Piktogramme hochladen", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

datum = datetime.today().strftime('%d.%m.%Y')

# PDF-Erstellung mit Vorlage und Koordinaten
def create_pdf(path, template_path):
    pdf = FPDF()
    pdf.add_page()

    # Hintergrundbild (PNG der DGUV-Vorlage)
    if os.path.exists(template_path):
        pdf.image(template_path, x=0, y=0, w=210, h=297)  # DIN A4

    pdf.set_font("Arial", size=10)

    # Kopfbereich
    pdf.set_xy(25, 32)
    pdf.multi_cell(0, 5, f"{firmenname} - Betriebsanweisung Nr. {nr}")

    pdf.set_xy(25, 42)
    pdf.multi_cell(0, 5, f"Arbeitsbereich: {arbeitsbereich}\nArbeitsplatz: {arbeitsplatz}\nTätigkeit: {taetigkeit}")

    # Gefahren
    pdf.set_xy(25, 80)
    pdf.multi_cell(160, 6, gefahren)

    # Schutzmaßnahmen
    pdf.set_xy(25, 120)
    pdf.multi_cell(160, 6, schutz)

    # Verhalten im Gefahrfall
    pdf.set_xy(25, 160)
    pdf.multi_cell(160, 6, verhalten_gefahrfall)

    # Erste Hilfe
    pdf.set_xy(25, 200)
    pdf.multi_cell(160, 6, erste_hilfe)

    # Entsorgung
    pdf.set_xy(25, 240)
    pdf.multi_cell(160, 6, entsorgung)

    # Piktogramme unten einfügen
    if pictogramme:
        x_pos = 25
        y_pos = 270
        for pict in pictogramme:
            tmp_path = os.path.join("/tmp", pict.name)
            with open(tmp_path, "wb") as f:
                f.write(pict.getbuffer())
            pdf.image(tmp_path, x=x_pos, y=y_pos, w=20)
            x_pos += 25

    pdf.set_xy(25, 285)
    pdf.cell(0, 5, f"Datum: {datum}   Unterschrift: ___________________", ln=True)

    pdf.output(path)

# DOCX-Erstellung (einfach gehalten)
def create_docx(path):
    doc = Document()
    doc.add_heading(f"{firmenname} - BETRIEBSANWEISUNG Nr. {nr}", 0)
    doc.add_paragraph(f"Arbeitsbereich: {arbeitsbereich}\nArbeitsplatz: {arbeitsplatz}\nTätigkeit: {taetigkeit}")

    sections = {
        "Gefahren für Mensch und Umwelt": gefahren,
        "Schutzmaßnahmen und Verhaltensregeln": schutz,
        "Verhalten im Gefahrfall": verhalten_gefahrfall,
        "Erste Hilfe": erste_hilfe,
        "Sachgerechte Entsorgung": entsorgung
    }

    for title, content in sections.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

    doc.add_paragraph(f"Datum: {datum}   Unterschrift: ___________________")
    doc.save(path)

# Datei-Upload für Vorlage
st.subheader("Vorlage hochladen (PNG der DGUV-Vorlage)")
template_file = st.file_uploader("Vorlage hochladen", type=["png"])

# Buttons für Export
if st.button("📥 PDF erstellen"):
    if template_file:
        template_path = os.path.join("/tmp", template_file.name)
        with open(template_path, "wb") as f:
            f.write(template_file.getbuffer())
    else:
        template_path = ""

    pdf_path = "/tmp/betriebsanweisung.pdf"
    create_pdf(pdf_path, template_path)
    with open(pdf_path, "rb") as f:
        st.download_button("PDF herunterladen", f, file_name="betriebsanweisung.pdf")

if st.button("📥 DOCX erstellen"):
    docx_path = "/tmp/betriebsanweisung.docx"
    create_docx(docx_path)
    with open(docx_path, "rb") as f:
        st.download_button("Word-Datei herunterladen", f, file_name="betriebsanweisung.docx")
